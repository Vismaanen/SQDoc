"""
Author		: paradowski.michal@outlook.com
Description	: main executable of SQDoc - MSSQL database documentation script
Updates:

* 2024-08-02 - v1.0 - creation
* 2024-08-27 - v1.0 - code cleanup, adding content settings enabling predefined content inclusion / exclusion
                      from printed document
"""

# import generic libraries
import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class MyPrinter:
    """
    Class responsible for exporting data into *.docx file.
    """

    def __init__(self, _details, _core):
        """
        Initialize cass instance.
        """
        # content settings
        self._content = _core.doc_content
        self._properties = _core.doc_properties
        # content objects
        self._db_config = _details.db_config
        self._db_tables = _details.db_tables
        self._db_procedures = _details.db_procedures
        # utilities
        self._db_name = _core.db_name
        self._log = _core.log
        self._utils = _core.utils
        self._export = _core.export
        # document export
        self._print_document()

    def _print_document(self):
        """
        Attempt to export database information to *.docx file.
        """
        print(f"----------\n{self._utils.timestamp()} creating *.docx file")
        doc = Document()
        section = doc.sections[0]
        _logo_path = os.path.join(os.path.dirname(__file__), '.resources', 'doc_logo.png')

        # ---------------------------------------------------------------------------------
        # header and footer
        # ---------------------------------------------------------------------------------
        header = section.header
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(_logo_path, width=Inches(1.0))
        run.add_text(f" {self._db_name} database technical documentation")

        # Add page numbers to each section's footer
        for section in doc.sections:
            self._add_page_numbers(section)

        # ---------------------------------------------------------------------------------
        # main page
        # ---------------------------------------------------------------------------------
        doc.add_heading(f'{self._db_name} database technical documentation', 0)
        # document details table
        config = {'header': ['Document properties', ''], 'columns': [2, 2]}
        self._add_table(doc, config, self._properties)
        # next page
        doc.add_page_break()

        # ---------------------------------------------------------------------------------
        # table of contents
        # ---------------------------------------------------------------------------------
        doc.add_heading('Table of contents', level=1)
        toc_paragraph = doc.add_paragraph()
        toc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._add_toc(toc_paragraph)
        # next page
        doc.add_page_break()

        # ---------------------------------------------------------------------------------
        # document purpose
        # ---------------------------------------------------------------------------------
        doc.add_heading('1. Document purpose', level=1)
        doc.add_paragraph(f'Purpose of this document is to provide a technical overview on details and structure of '
                          f'{self._db_name} database. Document covers configuration of the database itself and '
                          f'properties of each included table such as:')
        doc.add_paragraph(f'table columns', style='List Bullet')
        doc.add_paragraph(f'column types', style='List Bullet')
        doc.add_paragraph(f'keys', style='List Bullet')
        doc.add_paragraph(f'Details related to table contents and / or volume are not part o this document due to '
                          f'their potentially sensitive nature.')
        # next page
        doc.add_page_break()

        # ---------------------------------------------------------------------------------
        # database details
        # ---------------------------------------------------------------------------------
        if self._content['db_configuration']:
            doc.add_heading(f'2. {self._db_name} database details', level=1)
            doc.add_paragraph(f'This section covers basic configuration details of database.')
            # for each subject - create content table
            paragraph = 1
            for scope in self._db_config:
                content = self._db_config[scope]
                doc.add_heading(f"2.{paragraph} {scope}", level=2)
                # create table
                # separate formatting for different scope subjects
                if scope == 'Configuration':
                    config = {'header': ['Configuration item', 'Value default', 'Value in use'], 'columns': [4, 1, 1]}
                    self._add_table(doc, config, content)
                if scope == "Scoped configuration":
                    config = {'header': ['Configuration item', 'Value'], 'columns': [4, 2]}
                    self._add_table(doc, config, content)
                paragraph += 1
            # next page
            doc.add_page_break()

        # ---------------------------------------------------------------------------------
        # per-table details
        # ---------------------------------------------------------------------------------
        if self._content['db_tables']:
            doc.add_heading(f'3. {self._db_name} tables', level=1)
            doc.add_paragraph(f'This section covers basic configuration details of database tables.')
            # for each subject - create content table
            paragraph = 1
            for table in self._db_tables:
                content = self._db_tables[table]
                doc.add_heading(f"3.{paragraph} {table}", level=2)

                # section - keys
                doc.add_heading(f"3.{paragraph}.1 Keys", level=3)
                if len(content['keys']) == 0:
                    doc.add_paragraph(f'No keys configured for this table.')
                else:
                    config = {'header': ['Key column name', 'Constraint name', 'Constraint type'], 'columns': [2, 2, 2]}
                    self._add_table(doc, config, content['keys'])

                # section - extended properties
                doc.add_heading(f"3.{paragraph}.2 Extended properties", level=3)
                if len(content['extended']) == 0:
                    doc.add_paragraph(f'No extended properties configured for this table.')
                else:
                    config = {'header': ['Property name', 'Property value'], 'columns': [2, 4]}
                    self._add_table(doc, config, content['extended'])

                # section - columns
                doc.add_heading(f"3.{paragraph}.3 Columns", level=3)
                if len(content['columns']) == 0:
                    doc.add_paragraph(f'No columns configured for this table')
                else:
                    config = {'header': ['Column name', 'Data type', 'Max length', 'Nullable'], 'columns': [2, 2, 1, 1]}
                    self._add_table(doc, config, content['columns'])
                paragraph += 1

            # next page
            doc.add_page_break()

        # ---------------------------------------------------------------------------------
        # stored procedure details
        # ---------------------------------------------------------------------------------
        if self._content['db_procedures']:
            doc.add_heading(f'3. {self._db_name} stored procedures', level=1)
            doc.add_paragraph(f'This section covers basic configuration details of configured stored procedures.')
            # for each subject - create content table
            paragraph = 1
            for procedure in self._db_procedures:
                content = self._db_procedures[procedure]
                doc.add_heading(f"3.{paragraph} {procedure}", level=2)

                # section - extended properties
                doc.add_heading(f"3.{paragraph}.1 Extended properties", level=3)
                if not content['extended']:
                    doc.add_paragraph(f'No extended properties configured for this procedure.')
                else:
                    config = {'header': ['Property name', 'Property value'], 'columns': [2, 4]}
                    self._add_table(doc, config, content['extended'])

                # section - stored procedure details
                doc.add_heading(f"3.{paragraph}.2 Details", level=3)
                if not content['info']:
                    doc.add_paragraph(f'No properties obtained for this procedure.')
                else:
                    config = {'header': ['Property name', 'Property value'], 'columns': [2, 4]}
                    self._add_table(doc, config, content['info'])
                paragraph += 1

        # Save the document
        path = f"{self._export}\\{self._db_name}_documentation.docx"
        print(f"{self._utils.timestamp()} Saving file: {path}")
        doc.save(path)

    @staticmethod
    def _add_toc(paragraph):
        """
        Build table of content.
        :param paragraph: current page paragraph object
        :type paragraph: doc.add_paragraph()
        :return: modified paragraph with table of content configuration
        :rtype: doc.add_paragraph()
        """
        run = paragraph.add_run()
        char = OxmlElement('w:fldChar')
        char.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-2" \\h \\z \\u'
        # ---
        char_ext_1 = OxmlElement('w:fldChar')
        char_ext_1.set(qn('w:fldCharType'), 'separate')
        char_ext_2 = OxmlElement('w:updateFields')
        char_ext_2.set(qn('w:val'), 'true')
        char_ext_3 = OxmlElement('w:fldChar')
        char_ext_3.set(qn('w:fldCharType'), 'end')
        # ---
        r_element = run._r
        r_element.append(char)
        r_element.append(instr)
        r_element.append(char_ext_1)
        r_element.append(char_ext_2)
        r_element.append(char_ext_3)
        return paragraph

    @staticmethod
    def _add_page_numbers(section):
        """
        Method to provide pages numeration.
        :param section: document section object
        """
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # paragraph for storing numbers
        run = paragraph.add_run()
        run.font.size = Pt(12)
        char = OxmlElement('w:fldChar')
        char.set(qn('w:fldCharType'), 'begin')
        # ---
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        # ---
        char_ext = OxmlElement('w:fldChar')
        char_ext.set(qn('w:fldCharType'), 'end')
        # ---
        run._r.append(char)
        run._r.append(instr)
        run._r.append(char_ext)

        footer_text = footer.add_paragraph()
        footer_text.add_run("SQDoc 1.0")
        footer_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def _add_table(doc, config, content):
        """
        Create table for provided data.
        :param doc: python-docx document object
        :param config: table configuration containing header column names and column dimensions list
        :param content: table content list
        :type doc: docx.Document()
        :type config: dict(str, any)
        :type content: list(any)
        :return: modified document object
        :rtype: docx.Document()
        """
        cols = len(config['columns'])
        table = doc.add_table(rows=1, cols=cols)
        table.autofit = False
        table.style = 'Light List Accent 1'
        theader = table.rows[0].cells
        # set header text and column dimensions
        for item_id in range(0, cols):
            theader[item_id].text = config['header'][item_id]
            theader[item_id].width = Inches(config['columns'][item_id])
        # write table content
        for item in content:
            row = table.add_row().cells
            for item_id in range(0, cols):
                row[item_id].text = item[item_id]
                row[item_id].width = Inches(config['columns'][item_id])
        return doc


def execute(_details, _core):
    """
    Carry document print task.
    :param _details: MSSQL database details dictionary
    :param _core: SQDoc script object
    :type _details: dict(str, any)
    :type _core: SQDoc()
    :raise exc: Unspecified fil creation exception
    """
    try:
        MyPrinter(_details, _core)
        print(f"{_core.utils.timestamp()} Document saved, all activities finished.")
    except Exception as exc:
        print(f"{_core.utils.timestamp()} Unspecified exception, check log for details. Exiting...")
        _core.log.warn(f"Unspecified *.docx file creation exception: {exc}")
        sys.exit(0)
